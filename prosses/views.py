# Django core tools
from django.shortcuts import render, redirect, get_object_or_404
from django.http import Http404, FileResponse
from django.core.files import File
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.conf import settings

# External library
from docx import Document

# Python built-in
import os

# Models
from .models import Fiche_descptive
from modules_studets.models import Affectation, Filers_modules, AffectationSequence, AffectationChapitre



def get_doc_path(filename):
    return os.path.join(settings.BASE_DIR, 'templates_docs', filename)


def replace_placeholders(doc_path, replacements):
    """
    Remplace les placeholders dans un fichier Word par les valeurs données.
    """
    doc = Document(doc_path)

    for key in ["cc", "EFCF", "th", "pratique"]:
        if key in replacements:
            replacements[key] = "X" if replacements[key] else ""

    for para in doc.paragraphs:
        for key, value in replacements.items():
            if f"<<{key}>>" in para.text:
                para.text = para.text.replace(f"<<{key}>>", value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in replacements.items():
                        if f"<<{key}>>" in para.text:
                            para.text = para.text.replace(f"<<{key}>>", value)

    return doc


def format_list(items):
    """
    Formate une liste d'éléments avec une numérotation.
    """
    return "\n".join(f"{i}. {item}" for i, item in enumerate(items, start=1))

@login_required(login_url='login')
def generate_file(request):
    if request.method != "POST":
        return redirect('dach_formater')

    user = request.user
    affectation_id = request.POST.get('affectation_id')
    module_id = request.POST.get('module_id')


    if not all([affectation_id, module_id]):
        messages.error(request, "Paramètres manquants")
        return redirect('dach_formater')

    affectation = get_object_or_404(Affectation, id=affectation_id, formateur=user)
    module = get_object_or_404(affectation.modules, id=module_id)
    model_dexaman = get_object_or_404(Filers_modules, id=module_id)

    # récupère les IDs des séquences cochées
    sequence_ids = request.POST.getlist('sequences[]')
    criteres = format_list(request.POST.getlist('critere[]'))
    porsontage = format_list(request.POST.getlist('pourcentage[]'))
    objectifs = format_list(request.POST.getlist('objectifs[]'))

    # récupère seulement les AffectationSequence sélectionnées
    affectation_sequences = AffectationSequence.objects.filter(
        id__in=sequence_ids,
        affectation=affectation,
        sequence__module=module
    ).select_related('sequence').order_by('sequence__order')

    # récupère les chapitres liés à ces séquences
    chapitres = AffectationChapitre.objects.filter(
        affectation_sequence__in=affectation_sequences
    ).select_related('chapitre').order_by('chapitre__order')

    # extrait les numéros des séquences sélectionnées
    sequence_orders = [str(seq.sequence.order) for seq in affectation_sequences]

    # fonction pour joindre avec 'et' بشكل لطيف
    def join_with_et(items):
        if not items:
            return ""
        if len(items) == 1:
            return items[0]
        return ", ".join(items[:-1]) + " et " + items[-1]

    # préparer les valeurs à remplacer
    replacements = {
        "annee_formation": affectation.filiere.academic_year.annee,
        "formateur": user.full_name,
        "filiere": affectation.filiere.title,
        "unite_formation": model_dexaman.titer_module,
        "objectif": str(model_dexaman.titer_module_pransipale),
        "conditions": format_list([chap.chapitre.titre for chap in chapitres]),
        "cc": False,
        "EFCF": False,
        "th": False,
        "pratique": False,
        "date_epreuve": request.POST.get('date_ducontrole', ""),
        "numero": request.POST.get('nemuro_controle', ""),
        "duree": request.POST.get('dure_ducontrole', ""),
        "sequences": join_with_et(sequence_orders),
        "criteres": criteres,
        "ponderation": porsontage,
        "explication": objectifs
    }

    # gérer les types
    controle_type = request.POST.get('controle_type')
    controle_mode = request.POST.get('controle_mode')

    if not controle_type or not controle_mode:
        messages.error(request, "Veuillez sélectionner un type ET un mode de contrôle")
        return redirect('dach_formater')

    replacements.update({
        'cc': controle_type == 'cc',
        'EFCF': controle_type == 'efcf',
        'th': controle_mode == 'th',
        'pratique': controle_mode == 'pr'
    })

    type_mapping = {
        ('cc', 'th'): 'Contrôle Continu-Théorique',
        ('cc', 'pr'): 'Contrôle Continu-Pratique',
        ('efcf', 'th'): 'EFCF-Théorique',
        ('efcf', 'pr'): 'EFCF-Pratique',
    }
    tipe = type_mapping.get((controle_type, controle_mode), 'Type non spécifié')

    # génération du fichier Word
    doc_path = get_doc_path('template.docx')
    doc = replace_placeholders(doc_path, replacements)

    output_filename = f"{user.full_name}_{affectation.filiere.title}_{model_dexaman.titer_module}_fiche.docx"
    output_path = os.path.join(settings.BASE_DIR, 'templates_docs', 'output', output_filename)
    doc.save(output_path)

    with open(output_path, 'rb') as f:
        fiche_descptive = Fiche_descptive(
            title=f"{model_dexaman.titer_module} {user.full_name}",
            type=tipe,
            user=user,
            module=model_dexaman,
            filer=affectation.filiere
        )
        fiche_descptive.file.save(output_filename, File(f), save=True)

    if os.path.exists(output_path):
        os.remove(output_path)

    messages.success(request, "Fiche générée avec succès")
    return redirect('dach_formater')


def download_fiche_file(request, fiche_id):
    fiche = get_object_or_404(Fiche_descptive, pk=fiche_id)

    if fiche.file:
        try:
            return FileResponse(fiche.file.open('rb'), as_attachment=True, filename=fiche.file.name)
        except FileNotFoundError:
            raise Http404("Le fichier n'existe pas.")
    else:
        raise Http404("Aucun fichier lié à cette fiche.")


def dawnload_file(request, id):
    fiche = get_object_or_404(Fiche_descptive, id=id)
    if fiche.file:
        try:
            return FileResponse(fiche.file.open('rb'), as_attachment=True, filename=fiche.file.name)
        except FileNotFoundError:
            raise Http404("Le fichier n'existe pas.")
    else:
        raise Http404("Le fichier n'existe pas.")